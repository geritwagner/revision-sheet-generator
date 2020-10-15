#! /usr/bin/env python
# -*- coding: utf-8 -*-

import os
import argparse
import docx
from docx.shared import Cm
import pylatex
from pytablewriter import MarkdownTableWriter

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

def generate_word_revision_sheet(filepath, result_path, lines, starting_item):
    document = docx.Document()

    document.add_heading('Revision sheet', 0)

    table = document.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nr.'
    hdr_cells[1].text = 'Comment'
    hdr_cells[2].text = 'How the comment is addressed'
    for line in lines:
        if not line.strip():
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = str(starting_item)
        row_cells[1].text = line.replace('\\newline', '')
        row_cells[2].text = ''
        starting_item += 1
    
    set_column_width(table.columns[0], Cm(1.5))
    set_column_width(table.columns[1], Cm(10))
    set_column_width(table.columns[2], Cm(6))
    if result_path:
        document.save(result_path)
    else:
        document.save(filepath[:-4] + '_revision_sheet.docx')
    return

def generate_tex_revision_sheet(filepath, result_path,lines, starting_item):
    
    geometry_options = {"tmargin": "2.54cm", "lmargin": "2.54cm"}
    doc = pylatex.Document(geometry_options=geometry_options)

    with doc.create(pylatex.Section('Revision sheet')):

        with doc.create(pylatex.LongTable('|r|p{8cm}|p{8cm}|')) as table:
            table.add_hline()
            table.add_row(('Nr.', 'Comment', 'How the comment is addressed'))
            table.add_hline()
            for line in lines:
                if not line.strip():
                    continue
                table.add_row((starting_item, line.replace('\\newline', '').replace('â€¯',''), ''))
                table.add_hline()
                starting_item += 1
    
    if result_path:
        doc.generate_pdf(result_path, clean_tex=False)
    else:
        doc.generate_pdf(filepath[:-4] + '_revision_sheet.tex', clean_tex=False)

    return


def generate_md_revision_sheet(filepath, result_path,lines, starting_item):
    
    writer = MarkdownTableWriter()
    if result_path:
        result_path = result_path
    else:
        result_path = result_path + "revision_sheet.md"
    writer.headers = ["Nr.", "Comment", "How the comment is addressed"]
    
    comment_list = []
    for line in lines:
        if not line.strip():
            continue
        comment_list.append(line.replace('\\newline', ''))

    matrix = []
    for i in range(len(comment_list)):
        matrix.append([starting_item, comment_list[i], ""])        
        starting_item += 1
        print(matrix)

    writer.value_matrix = matrix

    writer.write_table()
    
    with open(result_path, "w") as f:
        writer.stream = f
        writer.write_table()
    
    return

def load_file(filepath):
    comment_file = open(filepath, 'r') 
    lines = comment_file.readlines()
    revised_lines = []
    temp = ''
    for line in lines:
        if '\\newline' in line or ('Reviewer' in line and len(line.strip()) < 15):
            temp += line
            continue
        if '' == temp:
            revised_lines.append(line.rstrip())
        else:
            revised_lines.append(temp + line.rstrip())
            temp = ''
    if '' != temp:
        revised_lines.append(temp.rstrip())
    lines = revised_lines  
    return lines

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description = "Revision-sheet generator")
    parser.add_argument("--input", default=None, help="path to the review text file") 
    parser.add_argument("--format", default='w', help="format of the output document , w for word (default) or t for tex or m for markdown") 
    parser.add_argument("--output", default=None, help="path to the file where to put the results (optional)") 
    parser.add_argument("--i", default=1, help="start of comment numbering (optional)") 

    args = parser.parse_args()
    filepath = args.input
    output_format = args.format
    result_path = args.output
    if not result_path:
        result_path = 'revision_sheet'
    starting_item = int(args.i)
    
    assert output_format in ['t', 'w', 'm']

    if 'w' == output_format:
        if result_path:
            if '.doc' != result_path[:-4] or '.docx' != result_path[:-5]:
                result_path += '.doc'
    if 't' == output_format:
        if result_path:
            if '.tex' != result_path[:-4]:
                result_path += '.tex'
    if 'm' == output_format:
        if result_path:
            if '.md' != result_path[:-4]:
                result_path += '.md'
    
    assert filepath[-4:] == '.txt'
    assert isinstance(starting_item, int)
    if result_path:
        assert not os.path.exists(result_path)
    
    lines = load_file(filepath)    

    if 'w' == output_format:
        generate_word_revision_sheet(filepath, result_path, lines, starting_item)
    if 't' == output_format:
        generate_tex_revision_sheet(filepath, result_path, lines, starting_item)
    if 'm' == output_format:
        generate_md_revision_sheet(filepath, result_path, lines, starting_item)
        
